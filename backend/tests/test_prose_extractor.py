from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.modules.prose_extractor import extract_prose


def _add_bold_centered(doc: Document, text: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    return para


def test_explicit_heading_reclassifies_prior_heuristic_title_page_headings():
    doc = Document()
    _add_bold_centered(doc, "Organizational Change Implementation Plan")
    doc.add_paragraph("Dihan Zhang")
    doc.add_paragraph("University of Southern California")
    doc.add_paragraph("Introduction to the Problem", style="Heading 1")
    doc.add_paragraph(
        "This paragraph introduces the study context. It establishes the problem."
    )

    paragraphs = extract_prose(doc)

    assert paragraphs[0].heading_level == 0
    assert paragraphs[1].heading_level == 0
    assert paragraphs[2].heading_level == 0
    assert paragraphs[3].heading_level == 1
    assert paragraphs[4].heading_level is None


def test_all_manual_document_preserves_first_section_heading_after_preamble():
    doc = Document()
    _add_bold_centered(doc, "Organizational Change Implementation Plan")
    doc.add_paragraph("Dihan Zhang")
    doc.add_paragraph("University of Southern California")
    _add_bold_centered(doc, "Introduction to the Problem")
    doc.add_paragraph(
        "This paragraph introduces the study context. It establishes the problem."
    )

    paragraphs = extract_prose(doc)

    assert paragraphs[0].heading_level == 0
    assert paragraphs[1].heading_level == 0
    assert paragraphs[2].heading_level == 0
    assert paragraphs[3].heading_level == 1
    assert paragraphs[4].heading_level is None


def test_table_label_and_title_are_layout_elements():
    doc = Document()
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Demographic Characteristics of Participants by Program Type")
    doc.add_paragraph("This paragraph discusses the table. It explains the pattern.")

    paragraphs = extract_prose(doc)

    assert paragraphs[0].heading_level == 0
    assert paragraphs[1].heading_level == 0
    assert paragraphs[2].heading_level is None
    assert paragraphs[2].paragraph_number_on_page == 1


def test_table_label_and_title_do_not_increment_paragraph_number():
    doc = Document()
    doc.add_paragraph("This paragraph introduces the section. It has enough detail.")
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Demographic Characteristics of Participants by Program Type")
    doc.add_paragraph("This paragraph follows the table. It should be paragraph two.")

    paragraphs = extract_prose(doc)

    assert paragraphs[0].paragraph_number_on_page == 1
    assert paragraphs[3].paragraph_number_on_page == 2


def test_headings_and_list_paragraphs_do_not_increment_paragraph_number():
    doc = Document()
    doc.add_paragraph("Implementation Plan", style="Heading 1")
    doc.add_paragraph("This paragraph introduces the section. It has enough detail.")
    doc.add_paragraph("A layout list item", style="List Paragraph")
    doc.add_paragraph("This paragraph follows the layout line. It should be paragraph two.")

    paragraphs = extract_prose(doc)

    assert paragraphs[1].paragraph_number_on_page == 1
    assert paragraphs[3].paragraph_number_on_page == 2


def test_numbered_list_metadata_does_not_increment_paragraph_number():
    doc = Document()
    doc.add_paragraph("This paragraph introduces the section. It has enough detail.")
    list_para = doc.add_paragraph("This is a numbered list item with Normal style.")
    pPr = list_para._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(num_id)
    pPr.append(numPr)
    doc.add_paragraph("This paragraph follows the numbered item. It should be paragraph two.")

    paragraphs = extract_prose(doc)

    assert paragraphs[1].heading_level == 0
    assert paragraphs[1].is_list_item
    assert paragraphs[0].paragraph_number_on_page == 1
    assert paragraphs[2].paragraph_number_on_page == 2


def test_bold_italic_manual_heading_is_level_three():
    doc = Document()
    doc.add_paragraph("Main Section", style="Heading 1")
    para = doc.add_paragraph()
    run = para.add_run("Pilot Program Data")
    run.bold = True
    run.italic = True

    paragraphs = extract_prose(doc)

    assert paragraphs[1].heading_level == 3


def test_bold_sentence_with_terminal_punctuation_is_not_heading():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("The program included three stages: (1) interviews, (2) workshops, and (3) tasks.")
    run.bold = True

    paragraphs = extract_prose(doc)

    assert paragraphs[0].heading_level is None


def test_page_break_resets_paragraph_number_on_page():
    doc = Document()
    doc.add_paragraph("This is the first paragraph. It is on page one.")
    doc.add_page_break()
    doc.add_paragraph("This is the second paragraph. It is on page two.")

    paragraphs = extract_prose(doc)

    assert paragraphs[0].page_number == 1
    assert paragraphs[0].paragraph_number_on_page == 1
    assert paragraphs[1].page_number == 2
    assert paragraphs[1].paragraph_number_on_page == 1

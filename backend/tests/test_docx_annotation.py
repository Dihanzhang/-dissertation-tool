from docx import Document

from app.main import _annotated_rows, _highlight_runs
from app.modules.module2_apa_checker import Finding, Severity
from app.modules.module3_citation_matcher import CitationMatchResult


def test_highlight_runs_splits_single_run_to_target_phrase_only():
    doc = Document()
    para = doc.add_paragraph("The initiative is delivered by four coordinated teams.")

    anchor_runs = _highlight_runs(para, "is delivered", "SUGGESTION")

    assert para.text == "The initiative is delivered by four coordinated teams."
    assert [run.text for run in anchor_runs] == ["is delivered"]

    highlighted = [run.text for run in para.runs if run.font.highlight_color is not None]
    assert highlighted == ["is delivered"]


def test_highlight_runs_falls_back_to_one_sentence_not_whole_paragraph():
    doc = Document()
    para = doc.add_paragraph(
        "The first sentence should be marked. The second sentence should remain plain."
    )

    anchor_runs = _highlight_runs(para, "missing target", "WARNING")

    assert para.text == "The first sentence should be marked. The second sentence should remain plain."
    assert [run.text for run in anchor_runs] == ["The first sentence should be marked."]

    highlighted = [run.text for run in para.runs if run.font.highlight_color is not None]
    assert highlighted == ["The first sentence should be marked."]


def test_annotated_rows_deduplicates_same_rule_target_and_paragraph():
    findings = [
        Finding(
            rule_id="STY001",
            severity=Severity.SUGGESTION,
            paragraph_index=4,
            message="Passive voice",
            suggested_fix="",
            excerpt="are constrained",
        ),
        Finding(
            rule_id="STY001",
            severity=Severity.SUGGESTION,
            paragraph_index=4,
            message="APA §4.13: Passive voice detected ('are constrained'). Prefer active voice.",
            suggested_fix="Recast as active voice.",
            excerpt="are constrained",
        ),
    ]
    citation_result = CitationMatchResult([], [], [], [], [], [], [], [])

    rows = _annotated_rows(findings, citation_result)

    assert len(rows) == 1
    assert rows[0]["target"] == "are constrained"
    assert "APA §4.13" in rows[0]["message"]


def test_annotated_rows_targets_full_multiple_citation_parenthetical():
    citation = "(Lazzara et al., 2021; Hughes et al., 2018)"
    findings = [
        Finding(
            rule_id="CIT004",
            severity=Severity.WARNING,
            paragraph_index=8,
            message=(
                "APA §8.19: Multiple citations in one set of parentheses must be ordered "
                "alphabetically by first author's surname. "
                f"Found: '{citation}'."
            ),
            suggested_fix="Reorder citations alphabetically by first author surname",
            excerpt=citation,
        ),
    ]
    citation_result = CitationMatchResult([], [], [], [], [], [], [], [])

    rows = _annotated_rows(findings, citation_result)

    assert rows[0]["target"] == citation


def test_annotated_rows_ignores_possessive_apostrophe_when_targeting():
    findings = [
        Finding(
            rule_id="MEC026",
            severity=Severity.WARNING,
            paragraph_index=4,
            message=(
                "APA §6.15: Generic job titles are lowercase when they do not immediately "
                "precede a person's name. Use lowercase for 'Chief Executive Officer'."
            ),
            suggested_fix="Use lowercase: 'chief executive officer'",
            excerpt="shaped uptake. The Chief Executive Officer of the firm sponsor",
        ),
    ]
    citation_result = CitationMatchResult([], [], [], [], [], [], [], [])

    rows = _annotated_rows(findings, citation_result)

    assert rows[0]["target"] == "Chief Executive Officer"

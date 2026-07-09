"""
FastAPI backend — Build Step 2.

Endpoints:
  GET  /api/health
  POST /api/check/text        — Module 2 + 3 (free, deterministic)
  POST /api/check/docx        — Module 2 + 3 on uploaded .docx (free)
  POST /api/review/estimate   — word count + credit cost estimate (no LLM call)
  POST /api/review            — Module 1 (LLM, metered via credits)
"""

from __future__ import annotations

import math
import os
import re
import tempfile
from pathlib import Path

# Load .env in development (no-op if python-dotenv not installed or file absent)
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent.parent / ".env")
except ImportError:
    pass

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from starlette.background import BackgroundTask

from .modules.prose_extractor import ProseParagraph, extract_prose, QUOTE_MASK
from .modules.module2_apa_checker import check_paragraphs, Finding, _load_config
from .modules.module3_citation_matcher import (
    CitationMatchResult,
    run_citation_check,
    run_citation_check_paragraphs,
)
from .modules.module1_editor import run_module1, EditSuggestion
from .modules.text_splitter import count_words, split_into_chunks, estimate_chunks
from .modules import credits as credit_store

app = FastAPI(
    title="Dissertation APA 7 Review API",
    version="0.2.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],   # restrict in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_CONFIG_PATH = Path(__file__).parent / "config" / "prof_checklist.yaml"
WORD_CAP = int(os.getenv("WORD_CAP_PER_CREDIT", "5000"))
FREE_TRIAL_CAP = int(os.getenv("FREE_TRIAL_WORD_CAP", "3000"))
TEST_MODE = os.getenv("TEST_MODE", "false").lower() in ("true", "1", "yes")
RULE_SET_VERSION = "apa-qa-gap-f6d2709"


def _cfg():
    return _load_config(str(_CONFIG_PATH))


# ---------------------------------------------------------------------------
# Pydantic models — check endpoints (Modules 2 + 3)
# ---------------------------------------------------------------------------

class TextCheckRequest(BaseModel):
    body_text: str
    reference_text: str = ""
    levenshtein_threshold: int = 2


class FindingOut(BaseModel):
    rule_id: str
    severity: str
    paragraph_index: int
    message: str
    suggested_fix: str = ""
    autofixable: bool = False
    excerpt: str = ""
    location_hint: str = ""


class CheckResponse(BaseModel):
    apa_findings: list[FindingOut]
    missing_references: list[dict]
    uncited_references: list[dict]
    year_mismatches: list[dict]
    spelling_mismatches: list[dict]
    co_author_only_matches: list[dict]
    scope_warning: str = ""
    stats: dict


# ---------------------------------------------------------------------------
# Pydantic models — review endpoint (Module 1)
# ---------------------------------------------------------------------------

class ReviewRequest(BaseModel):
    body_text: str
    request_id: str
    user_id: str = "anonymous"
    tier: str = "paid"    # "free" | "paid"
    confirmed_oversized: bool = False  # user confirmed multi-credit split


class SuggestionOut(BaseModel):
    original: str
    revised: str
    reason: str
    edit_type: str        # "light" | "heavy"
    change_ratio: float


class ReviewResponse(BaseModel):
    status: str           # "ok" | "oversized_confirmation" | "no_credits" | "error"
    suggestions: list[SuggestionOut] = []
    # Oversized confirmation fields
    word_count: int = 0
    chunk_count: int = 0
    credits_required: int = 0
    credits_remaining: int = 0
    # Stats
    model_used: str = ""
    rejected_by_citation_lock: int = 0
    rejected_sentence_not_found: int = 0
    message: str = ""
    test_mode: bool = False


class EstimateRequest(BaseModel):
    body_text: str
    user_id: str = "anonymous"
    tier: str = "paid"


class EstimateResponse(BaseModel):
    word_count: int
    chunk_count: int
    credits_required: int
    credits_remaining: int
    word_cap_per_credit: int
    free_trial_cap: int
    test_mode: bool


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _finding_to_out(f: Finding) -> FindingOut:
    return FindingOut(
        rule_id=f.rule_id,
        severity=f.severity.value,
        paragraph_index=f.paragraph_index,
        message=f.message,
        suggested_fix=f.suggested_fix,
        autofixable=f.autofixable,
        excerpt=f.excerpt,
        location_hint=f.location_hint,
    )


def _suggestion_to_out(s: EditSuggestion) -> SuggestionOut:
    return SuggestionOut(
        original=s.original,
        revised=s.revised,
        reason=s.reason,
        edit_type=s.edit_type,
        change_ratio=s.change_ratio,
    )


def _paragraph_lookup(paragraphs: list[ProseParagraph]) -> dict[int, ProseParagraph]:
    return {p.index: p for p in paragraphs}


def _paragraph_order(para: ProseParagraph | None, fallback: int = 10**9) -> tuple[int, int, int]:
    if para is None:
        return (10**9, 10**9, fallback)
    return (
        getattr(para, "page_number", 1),
        getattr(para, "paragraph_number_on_page", para.index + 1),
        para.index,
    )


def _sort_findings(findings: list[Finding], paragraphs: list[ProseParagraph]) -> list[Finding]:
    para_by_index = _paragraph_lookup(paragraphs)
    return sorted(
        findings,
        key=lambda f: (*_paragraph_order(para_by_index.get(f.paragraph_index), f.paragraph_index), f.rule_id),
    )


def _decorate_and_sort_citation_issues(
    citation_result: CitationMatchResult,
    paragraphs: list[ProseParagraph],
) -> None:
    para_by_index = _paragraph_lookup(paragraphs)

    def decorate(issue: dict) -> dict:
        para = para_by_index.get(issue.get("paragraph_index"))
        if para is not None:
            issue["page_number"] = para.page_number
            issue["paragraph_number_on_page"] = para.paragraph_number_on_page
            anchor = para.raw_text[:60].strip()
            if len(para.raw_text) > 60:
                anchor += "..."
            issue["location_hint"] = (
                f'Page {para.page_number}, Para {para.paragraph_number_on_page} - "{anchor}"'
            )
        return issue

    def sort_key(issue: dict) -> tuple[int, int, int, int]:
        para = para_by_index.get(issue.get("paragraph_index"))
        page, para_on_page, doc_index = _paragraph_order(
            para,
            int(issue.get("line_index", 10**9)),
        )
        return (page, para_on_page, doc_index, int(issue.get("line_index", 0)))

    for attr in (
        "missing_references",
        "year_mismatches",
        "spelling_mismatches",
        "co_author_only_matches",
        "uncited_references",
    ):
        issues = [decorate(issue) for issue in getattr(citation_result, attr)]
        setattr(citation_result, attr, sorted(issues, key=sort_key))


_RULE_LABELS = {
    "PRF001": "Short paragraph",
    "REF010": "Publisher business designation",
    "STY001": "Passive voice",
    "MEC023": "First-line paragraph indent",
}


def _rule_label(rule_id: str, message: str) -> str:
    if rule_id in _RULE_LABELS:
        return _RULE_LABELS[rule_id]
    without_apa = re.sub(r"^APA\s+§[\d.]+:\s*", "", message)
    first_clause = re.split(r"[.:(-]", without_apa, maxsplit=1)[0].strip()
    return first_clause or rule_id


def _severity_definition(severity: str) -> str:
    definitions = {
        "ERROR": "Required fix: likely APA/reference inconsistency that should be corrected before submission.",
        "WARNING": "Review carefully: likely APA issue or formatting problem; fix unless your program requires otherwise.",
        "SUGGESTION": "Optional improvement: writing/style recommendation; accept only if it improves clarity and meaning.",
        "INFO": "Informational/program preference: check against your dissertation chair or school template.",
    }
    return definitions.get(severity.upper(), "Review this item and decide whether a change is needed.")


def _sentence_spans(text: str) -> list[tuple[int, int]]:
    spans = [(m.start(), m.end()) for m in re.finditer(r'[^.!?]+[.!?]?', text)]
    return [(s, e) for s, e in spans if text[s:e].strip()]


def _token_set(text: str) -> set[str]:
    return {t.lower() for t in re.findall(r"[A-Za-z0-9']+", text) if len(t) > 2}


def _best_sentence_span(text: str, target: str = "") -> tuple[int, int] | None:
    spans = _sentence_spans(text)
    if not spans:
        return (0, min(len(text), 120)) if text else None
    target_tokens = _token_set(target)
    if target_tokens:
        best = max(
            spans,
            key=lambda span: len(_token_set(text[span[0]:span[1]]) & target_tokens),
        )
        if len(_token_set(text[best[0]:best[1]]) & target_tokens) > 0:
            return best
    return spans[0]


def _target_span(text: str, target: str = "") -> tuple[int, int] | None:
    target = (target or "").strip()
    if target:
        start = text.find(target)
        if start >= 0:
            return (start, start + len(target))
        start = text.lower().find(target.lower())
        if start >= 0:
            return (start, start + len(target))
    return _best_sentence_span(text, target)


def _insert_run_after(run, text: str):
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.run import Run

    new_r = OxmlElement("w:r")
    rPr = run._r.find(qn("w:rPr"))
    if rPr is not None:
        new_r.append(deepcopy(rPr))
    t = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        t.set(qn("xml:space"), "preserve")
    t.text = text
    new_r.append(t)
    run._r.addnext(new_r)
    return Run(new_r, run._parent)


def _style_anchor_run(run, highlight_color, font_color) -> None:
    run.font.highlight_color = highlight_color
    run.font.color.rgb = font_color
    run.bold = True


def _highlight_runs(paragraph, target: str = "", severity: str = "WARNING") -> list:
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import RGBColor

    colors = {
        "ERROR": WD_COLOR_INDEX.RED,
        "WARNING": WD_COLOR_INDEX.YELLOW,
        "SUGGESTION": WD_COLOR_INDEX.TURQUOISE,
        "INFO": WD_COLOR_INDEX.BRIGHT_GREEN,
    }
    font_colors = {
        "ERROR": RGBColor(192, 0, 0),
        "WARNING": RGBColor(156, 101, 0),
        "SUGGESTION": RGBColor(0, 102, 204),
        "INFO": RGBColor(0, 128, 0),
    }
    highlight_color = colors.get(severity.upper(), WD_COLOR_INDEX.YELLOW)
    font_color = font_colors.get(severity.upper(), RGBColor(156, 101, 0))

    runs = [run for run in paragraph.runs if run.text]
    if not runs:
        return []

    para_text = paragraph.text or ""
    span = _target_span(para_text, target)
    if span is None:
        return []
    start, end = span

    highlighted = []
    cursor = 0
    for run in runs:
        original = run.text
        run_start = cursor
        run_end = cursor + len(original)
        if run_end > start and run_start < end:
            local_start = max(0, start - run_start)
            local_end = min(len(original), end - run_start)
            before = original[:local_start]
            anchor = original[local_start:local_end]
            after = original[local_end:]
            if local_start == 0 and local_end == len(original):
                _style_anchor_run(run, highlight_color, font_color)
                highlighted.append(run)
            elif local_start == 0:
                run.text = anchor
                if after:
                    _insert_run_after(run, after)
                _style_anchor_run(run, highlight_color, font_color)
                highlighted.append(run)
            else:
                run.text = before
                if after:
                    _insert_run_after(run, after)
                if anchor:
                    anchor_run = _insert_run_after(run, anchor)
                    _style_anchor_run(anchor_run, highlight_color, font_color)
                    highlighted.append(anchor_run)
        cursor = run_end
    return highlighted


def _comment_text(row: dict) -> str:
    message = str(row["message"])
    label = str(row["label"])
    if label and message.lower().startswith(label.lower()):
        message = message[len(label):].lstrip(": -")
    parts = [
        f'{row["kind"]}: {row["label"]}',
        f'Severity: {row["severity"]}',
        _severity_definition(str(row["severity"])),
        message,
    ]
    if row.get("suggested_fix"):
        parts.append(f'Suggested fix: {row["suggested_fix"]}')
    return "\n".join(parts)


def _finding_target(finding: Finding) -> str:
    if finding.rule_id == "CIT004" and finding.excerpt:
        return finding.excerpt
    found = re.search(r"Found:\s*'([^']+)'", finding.message)
    if found:
        return found.group(1)
    quoted = re.search(r"'([^']{2,80})'", finding.message)
    if quoted:
        return quoted.group(1)
    return finding.excerpt or ""


def _citation_issue_groups(citation_result: CitationMatchResult) -> list[tuple[str, dict]]:
    groups: list[tuple[str, dict]] = []
    for issue in citation_result.spelling_mismatches:
        groups.append(("Possible spelling mismatch", issue))
    for issue in citation_result.year_mismatches:
        groups.append(("Year mismatch", issue))
    for issue in citation_result.missing_references:
        label = "Missing reference" if issue.get("severity") == "error" else "Possible missing reference"
        groups.append((label, issue))
    for issue in citation_result.co_author_only_matches:
        groups.append(("Co-author-only match", issue))
    for issue in citation_result.uncited_references:
        groups.append(("Reference not cited in text", issue))
    return groups


def _annotated_rows(
    apa_findings: list[Finding],
    citation_result: CitationMatchResult,
) -> list[dict]:
    rows: list[dict] = []
    for finding in apa_findings:
        rows.append({
            "kind": "APA",
            "label": _rule_label(finding.rule_id, finding.message),
            "severity": finding.severity.value.upper(),
            "paragraph_index": finding.paragraph_index,
            "location": finding.location_hint,
            "context": finding.excerpt,
            "message": finding.message,
            "suggested_fix": finding.suggested_fix,
            "target": _finding_target(finding),
        })

    for label, issue in _citation_issue_groups(citation_result):
        rows.append({
            "kind": "Citation",
            "label": label,
            "severity": str(issue.get("severity", "warning")).upper(),
            "paragraph_index": issue.get("paragraph_index"),
            "location": issue.get("location_hint", ""),
            "context": issue.get("citation") or issue.get("reference", ""),
            "message": issue.get("message", ""),
            "suggested_fix": "",
            "target": issue.get("citation", ""),
        })

    deduped: dict[tuple, dict] = {}
    for row in rows:
        key = (
            row.get("paragraph_index"),
            str(row.get("kind", "")).lower(),
            str(row.get("label", "")).lower(),
            str(row.get("severity", "")).lower(),
            re.sub(r"\s+", " ", str(row.get("target", "")).strip().lower()),
        )
        existing = deduped.get(key)
        if existing is None:
            deduped[key] = row
            continue
        existing_score = len(str(existing.get("message", ""))) + len(str(existing.get("suggested_fix", "")))
        row_score = len(str(row.get("message", ""))) + len(str(row.get("suggested_fix", "")))
        if row_score > existing_score:
            deduped[key] = row

    return sorted(
        deduped.values(),
        key=lambda row: (
            row.get("paragraph_index") is None,
            int(row.get("paragraph_index") or 10**9),
            row["kind"],
            row["label"],
        ),
    )


def _annotate_docx(doc_path: str, output_path: str, rows: list[dict]) -> None:
    from docx import Document

    doc = Document(doc_path)
    for row in rows:
        para_index = row.get("paragraph_index")
        if not isinstance(para_index, int):
            continue
        if para_index < 0 or para_index >= len(doc.paragraphs):
            continue
        anchor_runs = _highlight_runs(
            doc.paragraphs[para_index],
            str(row.get("target") or ""),
            str(row.get("severity") or "WARNING"),
        )
        if anchor_runs:
            doc.add_comment(
                anchor_runs,
                text=_comment_text(row),
                author="Dissertation Review",
                initials="DR",
            )

    doc.save(output_path)


def _cleanup_paths(paths: list[str]) -> None:
    for path in paths:
        try:
            if path and os.path.exists(path):
                os.unlink(path)
        except OSError:
            pass


_SENTENCE_END = re.compile(r'[.!?]\s*$')
_LIST_ITEM = re.compile(r'^[\d]+[.)]\s|^[-•*–]\s')
_PAREN_YEAR_IN_TEXT = re.compile(r'\(\d{4}[a-z]?\)')
_TABLE_FIGURE_LABEL = re.compile(r'^(?:Table|Figure)\s+\d+\s*$', re.IGNORECASE)


def _is_likely_heading(text: str) -> bool:
    """
    Heuristic for pasted plain text (no .docx style info): detect heading-like lines.
    Prevents prose rules (P001, N001, etc.) from firing on titles and section headings.

    Matches:
    - ALL-CAPS lines up to 200 chars (common for top-level headings)
    - Short mixed-case lines (≤ 200 chars) with no terminal sentence punctuation
    Excludes list items and lines with parenthetical citation years.
    """
    stripped = text.strip()
    if not stripped:
        return False
    if '\n' in stripped:                        # multi-line → body paragraph
        return False
    if _LIST_ITEM.match(stripped):              # numbered/bulleted list item
        return False
    if _PAREN_YEAR_IN_TEXT.search(stripped):    # contains a citation year → body prose
        return False
    # ALL CAPS short line → heading/title label
    letters = [c for c in stripped if c.isalpha()]
    if letters and all(c.isupper() for c in letters) and len(stripped) <= 200:
        return True
    # Mixed-case: short with no terminal sentence punctuation
    if len(stripped) > 200:
        return False
    if _SENTENCE_END.search(stripped):          # ends with . ! ? → sentence, not heading
        return False
    return True


def _build_paragraphs(text: str) -> list[ProseParagraph]:
    """Build minimal ProseParagraph list from raw text (no .docx structure)."""
    paras = []
    para_on_page = 0
    previous_was_table_figure_label = False
    for i, raw in enumerate(text.split("\n\n")):
        raw = raw.strip()
        if not raw:
            continue
        masked = re.sub(r'“[^”]*?”|"[^"]*?"', QUOTE_MASK, raw)
        is_table_figure_label = bool(_TABLE_FIGURE_LABEL.match(raw))
        is_table_figure_title = (
            previous_was_table_figure_label
            and len(raw) <= 200
            and not _SENTENCE_END.search(raw)
            and not _PAREN_YEAR_IN_TEXT.search(raw)
        )
        if is_table_figure_label or is_table_figure_title:
            heading_level = 0
        else:
            heading_level = 1 if _is_likely_heading(raw) else None
        if heading_level is None:
            para_on_page += 1
        paras.append(ProseParagraph(
            index=i,
            style_name="Heading 1" if heading_level else "Normal",
            raw_text=raw,
            masked_text=masked,
            heading_level=heading_level,
            is_reference_entry=False,
            page_number=1,
            paragraph_number_on_page=para_on_page,
        ))
        previous_was_table_figure_label = is_table_figure_label
    return paras

def _empty_citation_result(note: str) -> CitationMatchResult:
    return CitationMatchResult(
        citations=[], references=[],
        missing_references=[], uncited_references=[],
        year_mismatches=[], spelling_mismatches=[],
        co_author_only_matches=[], formatting_issues=[],
        scope_warning=note,
    )


# ---------------------------------------------------------------------------
# Endpoints — Modules 2 + 3 (free)
# ---------------------------------------------------------------------------

@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "build_step": 2,
        "test_mode": TEST_MODE,
        "rule_set_version": RULE_SET_VERSION,
    }


@app.post("/api/check/text", response_model=CheckResponse)
async def check_text(req: TextCheckRequest):
    """Module 2 + 3 on plain text — free, no LLM."""
    cfg = _cfg()
    prose_cfg = cfg.get("prose_rules", {})
    heading_cfg = cfg.get("heading_rules", {})
    citation_cfg = cfg.get("citation_rules", {})
    threshold = req.levenshtein_threshold or citation_cfg.get("spelling_mismatch_threshold", 2)

    paragraphs = _build_paragraphs(req.body_text)
    apa_findings = _sort_findings(check_paragraphs(paragraphs, prose_cfg, heading_cfg), paragraphs)

    citation_result = (
        run_citation_check_paragraphs(paragraphs, req.reference_text, levenshtein_threshold=threshold)
        if req.reference_text
        else _empty_citation_result("No reference list provided — citation matching skipped.")
    )

    _decorate_and_sort_citation_issues(citation_result, paragraphs)

    return CheckResponse(
        apa_findings=[_finding_to_out(f) for f in apa_findings],
        missing_references=citation_result.missing_references,
        uncited_references=citation_result.uncited_references,
        year_mismatches=citation_result.year_mismatches,
        spelling_mismatches=citation_result.spelling_mismatches,
        co_author_only_matches=citation_result.co_author_only_matches,
        scope_warning=citation_result.scope_warning,
        stats={
            "paragraphs_checked": len(paragraphs),
            "apa_findings_count": len(apa_findings),
            "citations_found": len(citation_result.citations),
            "references_parsed": len(citation_result.references),
        },
    )


@app.post("/api/check/docx", response_model=CheckResponse)
async def check_docx(
    file: UploadFile = File(...),
    reference_text: str = Form(default=""),
    levenshtein_threshold: int = Form(default=2),
):
    """Module 2 + 3 on an uploaded .docx — free. Uploaded file deleted after processing."""
    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")

    cfg = _cfg()
    prose_cfg = cfg.get("prose_rules", {})
    heading_cfg = cfg.get("heading_rules", {})
    citation_cfg = cfg.get("citation_rules", {})
    threshold = levenshtein_threshold or citation_cfg.get("spelling_mismatch_threshold", 2)

    content = await file.read()
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        from docx import Document
        doc = Document(tmp_path)
        paragraphs = extract_prose(doc)
        apa_findings = _sort_findings(check_paragraphs(paragraphs, prose_cfg, heading_cfg), paragraphs)

        effective_ref = reference_text or "\n".join(
            p.raw_text for p in paragraphs if p.is_reference_entry
        )
        citation_result = (
            run_citation_check_paragraphs(paragraphs, effective_ref, levenshtein_threshold=threshold)
            if effective_ref
            else _empty_citation_result("No reference list found — citation matching skipped.")
        )
        _decorate_and_sort_citation_issues(citation_result, paragraphs)
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

    return CheckResponse(
        apa_findings=[_finding_to_out(f) for f in apa_findings],
        missing_references=citation_result.missing_references,
        uncited_references=citation_result.uncited_references,
        year_mismatches=citation_result.year_mismatches,
        spelling_mismatches=citation_result.spelling_mismatches,
        co_author_only_matches=citation_result.co_author_only_matches,
        scope_warning=citation_result.scope_warning,
        stats={
            "paragraphs_checked": len(paragraphs),
            "apa_findings_count": len(apa_findings),
            "citations_found": len(citation_result.citations),
            "references_parsed": len(citation_result.references),
        },
    )


# ---------------------------------------------------------------------------
# Endpoints — Module 1 (LLM, metered)
# ---------------------------------------------------------------------------

@app.post("/api/check/docx/annotated")
async def check_docx_annotated(
    file: UploadFile = File(...),
    reference_text: str = Form(default=""),
    levenshtein_threshold: int = Form(default=2),
):
    """Return a copy of the uploaded .docx with findings highlighted and summarized."""
    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported.")

    cfg = _cfg()
    prose_cfg = cfg.get("prose_rules", {})
    heading_cfg = cfg.get("heading_rules", {})
    citation_cfg = cfg.get("citation_rules", {})
    threshold = levenshtein_threshold or citation_cfg.get("spelling_mismatch_threshold", 2)

    content = await file.read()
    tmp_path = None
    out_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        from docx import Document
        doc = Document(tmp_path)
        paragraphs = extract_prose(doc)
        apa_findings = _sort_findings(check_paragraphs(paragraphs, prose_cfg, heading_cfg), paragraphs)

        effective_ref = reference_text or "\n".join(
            p.raw_text for p in paragraphs if p.is_reference_entry
        )
        citation_result = (
            run_citation_check_paragraphs(paragraphs, effective_ref, levenshtein_threshold=threshold)
            if effective_ref
            else _empty_citation_result("No reference list found - citation matching skipped.")
        )
        _decorate_and_sort_citation_issues(citation_result, paragraphs)

        rows = _annotated_rows(apa_findings, citation_result)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as out:
            out_path = out.name
        _annotate_docx(tmp_path, out_path, rows)

        original_stem = Path(file.filename).stem or "document"
        safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", original_stem).strip("._") or "document"
        download_name = f"{safe_stem}_reviewed.docx"

        return FileResponse(
            out_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=download_name,
            background=BackgroundTask(_cleanup_paths, [tmp_path, out_path]),
        )
    except Exception:
        _cleanup_paths([p for p in [tmp_path, out_path] if p])
        raise


@app.post("/api/review/estimate", response_model=EstimateResponse)
async def review_estimate(req: EstimateRequest):
    """Return word count and credit cost estimate. No LLM call — free."""
    word_count = count_words(req.body_text)
    cap = FREE_TRIAL_CAP if req.tier == "free" else WORD_CAP
    chunks = math.ceil(word_count / cap) if word_count > 0 else 1
    return EstimateResponse(
        word_count=word_count,
        chunk_count=chunks,
        credits_required=chunks,
        credits_remaining=credit_store.get_balance(req.user_id),
        word_cap_per_credit=cap,
        free_trial_cap=FREE_TRIAL_CAP,
        test_mode=TEST_MODE,
    )


@app.post("/api/review", response_model=ReviewResponse)
async def review(req: ReviewRequest):
    """
    Module 1 — LLM-based clarity/voice editing, metered by credits.

    Flow:
    1. Check word count. If oversized and not confirmed → return confirmation prompt.
    2. Check credits.
    3. Split into chunks if oversized and confirmed.
    4. Run Module 1 per chunk; citation-lock each result.
    5. Decrement credit per successful chunk only.
    6. Return merged suggestions. Text is not retained after response.
    """
    is_trial = req.tier == "free"
    word_count = count_words(req.body_text)
    cap = FREE_TRIAL_CAP if is_trial else WORD_CAP
    chunk_count = math.ceil(word_count / cap) if word_count > 0 else 1
    balance = credit_store.get_balance(req.user_id)

    # Oversized: require explicit confirmation before spending multiple credits
    if chunk_count > 1 and not req.confirmed_oversized:
        return ReviewResponse(
            status="oversized_confirmation",
            word_count=word_count,
            chunk_count=chunk_count,
            credits_required=chunk_count,
            credits_remaining=balance,
            message=(
                f"This submission is ~{word_count:,} words and will be split into "
                f"{chunk_count} sections, using {chunk_count} credits. "
                f"You have {balance} credits remaining. Confirm to proceed."
            ),
            test_mode=TEST_MODE,
        )

    # Check if enough credits for all chunks
    if not TEST_MODE and not is_trial and balance < chunk_count:
        return ReviewResponse(
            status="no_credits",
            word_count=word_count,
            chunk_count=chunk_count,
            credits_required=chunk_count,
            credits_remaining=balance,
            message=(
                f"This submission needs {chunk_count} credits but you have {balance}. "
                "Top up to continue."
            ),
            test_mode=TEST_MODE,
        )

    # Credit gate (single-chunk or confirmed oversized)
    status = credit_store.check(
        user_id=req.user_id,
        request_id=req.request_id,
        word_count=min(word_count, cap),  # per-chunk check
        is_trial=is_trial,
    )
    if not status.allowed and not TEST_MODE:
        return ReviewResponse(
            status="no_credits",
            credits_remaining=status.credits_remaining,
            message=_credit_reason_message(status.reason),
            test_mode=TEST_MODE,
        )

    # Split text into chunks
    chunks = split_into_chunks(req.body_text, word_cap=cap)

    # Resolve provider once
    try:
        from .modules.provider import get_provider
        provider = get_provider(tier=req.tier)
    except ValueError as e:
        raise HTTPException(status_code=503, detail=str(e))

    all_suggestions: list[SuggestionOut] = []
    total_rejected_lock = 0
    total_rejected_nf = 0
    model_used = ""
    completed_chunks = 0

    for i, chunk in enumerate(chunks):
        chunk_request_id = f"{req.request_id}_chunk{i}"

        try:
            result = run_module1(chunk, provider=provider, tier=req.tier)
        except Exception as e:
            # Partial failure: report how many chunks succeeded
            return ReviewResponse(
                status="error",
                suggestions=all_suggestions,
                message=(
                    f"Error on chunk {i + 1}/{len(chunks)}: {e}. "
                    f"{completed_chunks} chunk(s) processed; credits charged: {completed_chunks}."
                ),
                model_used=model_used,
                rejected_by_citation_lock=total_rejected_lock,
                rejected_sentence_not_found=total_rejected_nf,
                test_mode=TEST_MODE,
            )

        # Decrement credit only after successful chunk
        credit_store.commit(
            user_id=req.user_id,
            request_id=chunk_request_id,
            is_trial=is_trial,
        )
        completed_chunks += 1
        model_used = result.model_used
        total_rejected_lock += result.rejected_by_citation_lock
        total_rejected_nf += result.rejected_sentence_not_found
        all_suggestions.extend(_suggestion_to_out(s) for s in result.suggestions)

    # Text is processed and not retained after this point
    return ReviewResponse(
        status="ok",
        suggestions=all_suggestions,
        word_count=word_count,
        chunk_count=len(chunks),
        credits_required=len(chunks),
        credits_remaining=credit_store.get_balance(req.user_id),
        model_used=model_used,
        rejected_by_citation_lock=total_rejected_lock,
        rejected_sentence_not_found=total_rejected_nf,
        message=f"Review complete. {len(all_suggestions)} suggestion(s) found.",
        test_mode=TEST_MODE,
    )


def _credit_reason_message(reason: str) -> str:
    if reason == "trial_already_used":
        return "Free trial already used. Purchase to continue."
    if reason == "no_credits":
        return "Credits used — top up to continue. APA rule checking remains free."
    if reason == "duplicate_request_id":
        return "This request was already processed."
    return f"Credits unavailable: {reason}"

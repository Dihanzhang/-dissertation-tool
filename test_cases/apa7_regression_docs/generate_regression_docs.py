from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
DOCS = ROOT / "docs"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_manual_heading(doc: Document, text: str, *, bold: bool = True, italic: bool = False, centered: bool = False) -> None:
    para = doc.add_paragraph()
    if centered:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic


def add_reference(doc: Document, text: str, *, italic_segment: str | None = None) -> None:
    para = doc.add_paragraph()
    if italic_segment and italic_segment in text:
        before, rest = text.split(italic_segment, 1)
        para.add_run(before)
        run = para.add_run(italic_segment)
        run.italic = True
        para.add_run(rest)
    else:
        para.add_run(text)


def mark_as_word_list_item(para) -> None:
    p_pr = para._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)


def save(doc: Document, filename: str) -> None:
    doc.save(DOCS / filename)


def mechanics_doc() -> list[str]:
    doc = Document()
    add_heading(doc, "Mechanics, Numbers, and Punctuation", 1)
    doc.add_paragraph(
        "15 participants completed the survey. The analysis revealed 3 themes and nearly thirty percent reported fatigue. "
        "The interviews occurred over five weeks, and participants completed 2 5-point scales."
    )
    doc.add_paragraph(
        "The 1990's data file listed 12000 responses from consultants across regions. Identity \u2014 not skill \u2014 emerged as the barrier. "
        "The barrier - not the technology - shaped adoption."
    )
    doc.add_paragraph(
        "Participants were assigned to the 1st cohort. The result was p = 0.031, p < .05, and r = 0.45."
    )
    doc.add_paragraph(
        'One participant stated, "...I felt new again," and later wrote that "professional identity is provisional...it changes."'
    )
    save(doc, "01_mechanics_numbers_punctuation.docx")
    return [
        "MEC002", "MEC003", "MEC004", "MEC031", "MEC022", "MEC011",
        "MEC010", "MEC012", "MEC024", "MEC025", "MEC007", "MEC009",
        "MEC008", "MEC028", "MEC021",
    ]


def citations_doc() -> list[str]:
    doc = Document()
    add_heading(doc, "Citation Formatting and Matching", 1)
    doc.add_paragraph(
        "Prior studies support the pattern (Nguyen, 2022; Brown, 2019; Adams, 2021). "
        "The same stream is reversed elsewhere (Kim, 2021; Kim, 2019). "
        "The team cited (Smith and Jones, 2020), while Taylor & Green (2021) described a different pattern."
    )
    doc.add_paragraph(
        "Several teams cited (Smith, Jones, Brown, 2020), (Miller et al, 2022), and (Davis, nd). "
        'One participant said, "I felt unprepared" (Ibarra, 2019). '
        "The method was confirmed by the sponsor (K. Patel, personal communication)."
    )
    doc.add_paragraph(
        "The model was discussed repeatedly (Adams, 2021). Adoption was slower than expected (Adams, 2021). "
        "Training improved confidence (Adams, 2021)."
    )
    add_heading(doc, "References", 1)
    add_reference(doc, "Adams, R. (2021). Capability frameworks. Journal of Workplace Learning, 33(4), 245-261.")
    add_reference(doc, "Brown, L. (2019). Technology acceptance. Organization Studies, 40(7), 989-1010.")
    add_reference(doc, "Ibarra, H. (2018). Working identity. Harvard Business Review Press.")
    add_reference(doc, "Kim, S. (2019). Identity threat. Academy of Management Review, 44(2), 301-320.")
    add_reference(doc, "Kim, S. (2021). Expertise renegotiation. Human Relations, 74(9), 1401-1425.")
    add_reference(doc, "Nguyen, T. (2022). Psychological readiness. Journal of Change Management, 22(3), 178-196.")
    add_reference(doc, "Zhou, M. (2020). Adult development and technological change. Management Learning, 51(5), 511-529.")
    save(doc, "02_citations_and_matching.docx")
    return [
        "CIT001", "CIT002", "CIT003", "CIT004", "CIT005", "CIT006",
        "CIT007", "CIT014", "CIT020", "CIT019",
        "year_mismatches", "missing_references", "uncited_references",
    ]


def style_bias_doc() -> list[str]:
    doc = Document()
    add_heading(doc, "Style and Bias-Free Language", 1)
    doc.add_paragraph(
        "The survey was completed by participants and the interviews were conducted by the team. "
        "The reason is because the rollout was new, and in order to support staff, the researcher conducted interviews."
    )
    doc.add_paragraph(
        "Participants that completed the pilot described a workflow which improved over time. "
        "Based on the results, the study concludes adoption increased. This shows a lot of progress."
    )
    doc.add_paragraph(
        "It would appear that change was likely. While adoption improved, fatigue remained. "
        "Ibid. was used incorrectly, and participants didn't always agree."
    )
    doc.add_paragraph(
        "Before a consultant adopts the tool, his professional identity may shift. "
        "A practitioner noted that he or she must renegotiate expertise. "
        "Several firms underestimated the manpower required."
    )
    doc.add_paragraph(
        "The sample included elderly participants, wheelchair-bound clients, the disabled, chairman roles, mankind references, "
        "Caucasian labels, seniors, African-American category names, and Asian-American category names."
    )
    save(doc, "03_style_bias_language.docx")
    return [
        "STY001", "STY002", "STY003", "STY004", "STY005", "STY006",
        "STY007", "STY008", "STY009", "STY010", "STY012", "STY013",
        "STY014", "STY015", "BFL001", "BFL006", "BFL011", "BFL018",
        "BFL018A", "BFL026", "BFL026A", "BFL031", "BFL032", "BFL033",
    ]


def headings_lists_tables_doc() -> list[str]:
    doc = Document()
    add_heading(doc, "Heading and List Regression", 1)
    add_manual_heading(doc, "Random Level Three Heading", bold=True, italic=True)
    doc.add_paragraph("This paragraph follows a skipped heading level. It provides body text.")
    add_heading(doc, "Empty Section", 2)
    add_heading(doc, "Next Empty Section", 2)
    doc.add_paragraph(
        "The program included three stages: (1) diagnostic interviews, (2) workflow mapping, and (3) supervised experimentation."
    )
    for item in [
        "Understanding the psychology of adoption.",
        "practitioners map their own workflows",
        "To evaluate outcomes honestly;",
    ]:
        para = doc.add_paragraph(item, style="List Bullet")
        mark_as_word_list_item(para)
    doc.add_paragraph("The results are shown in the table below and in Table III.")
    doc.add_paragraph("Short.")
    save(doc, "04_headings_lists_tables.docx")
    return ["HED001", "HED002", "MEC030", "MEC032", "TBL001", "TBL002", "PRF001"]


def reference_doc() -> list[str]:
    doc = Document()
    add_heading(doc, "Reference Formatting", 1)
    doc.add_paragraph("This body paragraph cites Smith (2020). It also cites Patel (2025).")
    add_heading(doc, "References", 1)
    add_reference(doc, "Smith, J. (2020). Technology Acceptance In Knowledge-Intensive Firms. Organization Studies, 40 (7), 989-1010. DOI: 10.1234/abc")
    add_reference(doc, "Taylor, A. and Green, B. (no date). Training outcomes. Journal of Change, 10(2), 11-20. Retrieved from https://example.com.")
    add_reference(doc, "Miller, C., et al. (2021). Edited volume title (2nd edition). Academic Press Inc. New York.")
    add_reference(doc, "Morgan, A. (2021). Legacy DOI format. Journal Name, 3(2), 9-15. http://doi.org/10.1234/legacy")
    add_reference(doc, "Brown, L. (2022) Missing date period. Journal Name, 5(1), 1-4.")
    add_reference(doc, "Davis, R. (In Press). Future research. Journal Name, 5(1), 1-4.")
    add_reference(doc, "Evans, P. (2021). A translated title (Trans.).")
    add_reference(doc, "Garcia, M. (2021). Original note. (originally published 1998).")
    add_reference(doc, "Kim, S. (2021). When the tool knows more than you: Expertise renegotiation in consulting. Human Relations, 74(9), 1401-1425.")
    add_reference(doc, "Patel, K. (2025, March 4). Personal communication.")
    add_reference(doc, "Zhou, M. (2020). Adult development. Management Learning, 51(5), 511-529. ISBN 1234567890")
    save(doc, "05_reference_formatting.docx")
    return [
        "REF001", "REF002", "REF003", "REF005", "REF006", "REF007",
        "REF008", "REF009", "REF010", "REF011", "REF012", "REF013",
        "REF015", "REF016", "REF023", "REF024", "REF025", "REF026",
    ]


def false_positive_doc() -> list[str]:
    doc = Document()
    add_heading(doc, "False Positive Traps", 1)
    para = doc.add_paragraph(
        "Twenty-three percent of participants completed the task over a period of 5 weeks. "
        "The president made remarks, and President Biden signed the order. "
        "Brown (2019) places the results in Table 3."
    )
    para.paragraph_format.first_line_indent = Inches(0.5)
    doc.add_paragraph("Table 1")
    doc.add_paragraph("Demographic Characteristics of Participants by Program Type")
    add_heading(doc, "References", 1)
    add_reference(
        doc,
        "Brown, L. (2019). Technology acceptance in knowledge-intensive firms. Organization Studies, 40(7), 989-1010.",
        italic_segment="Organization Studies, 40",
    )
    save(doc, "06_false_positive_traps.docx")
    return ["should_not_flag_time_unit_5_weeks", "should_not_flag_president_title", "should_not_flag_table_label_title"]


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    for old_doc in DOCS.glob("*.docx"):
        old_doc.unlink()

    cases = [
        ("01_mechanics_numbers_punctuation.docx", mechanics_doc()),
        ("02_citations_and_matching.docx", citations_doc()),
        ("03_style_bias_language.docx", style_bias_doc()),
        ("04_headings_lists_tables.docx", headings_lists_tables_doc()),
        ("05_reference_formatting.docx", reference_doc()),
        ("06_false_positive_traps.docx", false_positive_doc()),
    ]

    manifest = {
        "description": "DOCX regression suite for the dissertation APA 7 review tool.",
        "scope": "Covers deterministic APA 7 rule families currently implemented by the application.",
        "cases": [
            {
                "file": filename,
                "expected_rule_coverage": coverage,
            }
            for filename, coverage in cases
        ],
    }
    (ROOT / "manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
